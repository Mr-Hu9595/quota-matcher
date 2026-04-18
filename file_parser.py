"""
智能文件解析器
功能：自动识别文件类型（Word/Excel），解析工程量清单
"""

import re
from pathlib import Path
from typing import List, Dict, Optional
from abc import ABC, abstractmethod

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

try:
    from docx import Document
except ImportError:
    Document = None

from column_identifier import ColumnIdentifier


class ParseResult:
    """解析结果封装"""

    def __init__(self, items: List[Dict], warnings: List[str] = None):
        """
        初始化解析结果

        Args:
            items: 解析出的项目列表
            warnings: 解析警告信息
        """
        self.items = items
        self.warnings = warnings or []

    def __len__(self):
        return len(self.items)

    def __iter__(self):
        return iter(self.items)


class BaseParser(ABC):
    """文件解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> ParseResult:
        """
        解析文件

        Args:
            file_path: 文件路径

        Returns:
            ParseResult: 解析结果
        """
        pass


class ExcelParser(BaseParser):
    """Excel文件解析器"""

    def __init__(self):
        self.column_identifier = ColumnIdentifier()

    def parse(self, file_path: str) -> ParseResult:
        """
        解析Excel文件

        Args:
            file_path: Excel文件路径

        Returns:
            ParseResult: 解析结果
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".xls":
            return self._parse_xls(file_path)
        elif ext in [".xlsx", ".xlsm"]:
            return self._parse_xlsx(file_path)
        else:
            raise ValueError(f"不支持的Excel格式: {ext}")

    def _is_skippable_sheet(self, sheet_name: str) -> bool:
        """
        判断是否跳过该工作表

        Args:
            sheet_name: 工作表名称

        Returns:
            bool: 是否跳过
        """
        skip_keywords = ["封面", "总表", "目录", "说明", "备注", "表一", "表二"]
        sheet_lower = sheet_name.lower()
        return any(kw in sheet_lower for kw in skip_keywords)

    def _parse_xls(self, file_path: str) -> ParseResult:
        """解析旧版.xls格式Excel文件"""
        if xlrd is None:
            raise ImportError("请安装 xlrd: pip install xlrd")

        warnings = []
        all_items = []

        wb = xlrd.open_workbook(file_path)

        for sheet_index in range(wb.nsheets):
            ws = wb.sheet_by_index(sheet_index)
            sheet_name = ws.name

            if self._is_skippable_sheet(sheet_name):
                continue

            items, sheet_warnings = self._parse_xls_sheet(ws, sheet_name)
            all_items.extend(items)
            warnings.extend(sheet_warnings)

        return ParseResult(all_items, warnings)

    def _parse_xls_sheet(self, ws, sheet_name: str) -> tuple:
        """解析.xls工作表"""
        items = []
        warnings = []

        if ws.nrows < 2:
            return items, [f"工作表 '{sheet_name}' 数据行少于2行，跳过"]

        # 获取所有行
        rows = []
        for row_index in range(ws.nrows):
            row = [ws.cell_value(row_index, col_index) for col_index in range(ws.ncols)]
            rows.append(row)

        # 查找表头行
        header_row_index = self._find_header_row(rows)

        if header_row_index == -1:
            return items, [f"工作表 '{sheet_name}' 未找到有效表头，跳过"]

        headers = [str(h) if h is not None else "" for h in rows[header_row_index]]
        column_map = ColumnIdentifier.find_column_index(headers)

        valid, errors = ColumnIdentifier.validate_mapping(column_map)
        if not valid:
            return items, [f"工作表 '{sheet_name}': {', '.join(errors)}"]

        # 解析数据行
        for i in range(header_row_index + 1, len(rows)):
            row = rows[i]

            if all(cell == "" or cell is None for cell in row):
                continue

            data = ColumnIdentifier.extract_columns(list(row), column_map)

            if data["name"] and data["quantity"]:
                item = self._create_item(data, sheet_name)
                if item:
                    items.append(item)

        return items, warnings

    def _parse_xlsx(self, file_path: str) -> ParseResult:
        """解析新版.xlsx格式Excel文件"""
        if openpyxl is None:
            raise ImportError("请安装 openpyxl: pip install openpyxl")

        warnings = []
        all_items = []

        wb = openpyxl.load_workbook(file_path, data_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            if self._is_skippable_sheet(sheet_name):
                continue

            items, sheet_warnings = self._parse_sheet(ws, sheet_name)
            all_items.extend(items)
            warnings.extend(sheet_warnings)

        wb.close()

        return ParseResult(all_items, warnings)

    def _parse_sheet(self, ws, sheet_name: str) -> tuple:
        """
        解析单个工作表

        Args:
            ws: 工作表对象
            sheet_name: 工作表名称

        Returns:
            Tuple[List[Dict], List[str]]: (项目列表, 警告列表)
        """
        items = []
        warnings = []

        # 获取所有行数据
        rows = list(ws.iter_rows(values_only=True))

        if len(rows) < 2:
            return items, [f"工作表 '{sheet_name}' 数据行少于2行，跳过"]

        # 尝试找到表头行
        header_row_index = self._find_header_row(rows)

        if header_row_index == -1:
            warnings.append(f"工作表 '{sheet_name}' 未找到有效表头，跳过")
            return items, warnings

        headers = [str(h) if h is not None else "" for h in rows[header_row_index]]

        # 识别关键列
        column_map = ColumnIdentifier.find_column_index(headers)

        valid, errors = ColumnIdentifier.validate_mapping(column_map)
        if not valid:
            warnings.append(f"工作表 '{sheet_name}': {', '.join(errors)}")
            return items, warnings

        # 解析数据行
        for i in range(header_row_index + 1, len(rows)):
            row = rows[i]

            # 跳过空行
            if all(cell is None or str(cell).strip() == "" for cell in row):
                continue

            # 提取各列值
            data = ColumnIdentifier.extract_columns(list(row), column_map)

            # 验证并提取有效数据
            if data["name"] and data["quantity"]:
                item = self._create_item(data, sheet_name)
                if item:
                    items.append(item)

        return items, warnings

    def _find_header_row(self, rows: List[tuple]) -> int:
        """
        查找表头行索引

        Args:
            rows: 所有行数据

        Returns:
            int: 表头行索引，未找到返回-1
        """
        for i, row in enumerate(rows[:10]):  # 只检查前10行
            row_str = [str(cell).lower() if cell else "" for cell in row]
            row_text = "".join(row_str)

            # 检查是否包含关键词
            if any(keyword in row_text for keyword in
                   ["名称", "项目", "设备", "材料", "数量", "工程量", "工作量"]):
                return i

        return -1

    def _create_item(self, data: Dict[str, str], sheet_name: str) -> Optional[Dict]:
        """
        创建标准化的项目数据

        Args:
            data: 提取的列数据
            sheet_name: 工作表名称

        Returns:
            Dict: 标准化项目数据
        """
        name = data["name"].strip()
        quantity_str = data["quantity"].strip()
        unit = data.get("unit", "").strip()

        # 解析数量（可能是数字或表达式）
        try:
            # 尝试直接转换
            quantity = float(quantity_str)
        except ValueError:
            # 尝试解析表达式，如 "1+2" 或 "≈3.5"
            clean_str = quantity_str.replace("≈", "").replace("约", "").strip()
            try:
                quantity = float(clean_str)
            except ValueError:
                # 无法解析，跳过
                return None

        return {
            "name": name,
            "quantity": quantity,
            "unit": unit or "项",
            "sheet": sheet_name,
            "original_unit": unit  # 保留原始单位用于单位转换判断
        }


class WordParser(BaseParser):
    """Word文件解析器"""

    # 正则模式：共计：数字单位 或 共计：数字 单位
    QIJEI_PATTERN = re.compile(r'共计[：:]\s*(\d+\.?\d*)\s*(\S+)')

    # 项目分割模式：按"共计："分割，每段是一个独立项目
    ITEM_SPLIT_PATTERN = re.compile(r'共计[：:]\s*\d+')

    # 从事由中提取汇总数据的正则
    REASON_SUMMARY_PATTERN = re.compile(
        r'([^\s\d]+(?:\d+[^\s\d]*)?)\s*(\d+(?:\.\d+)?)\s*(米|个|套|台|根|节|块|千克|kg|吨|t)'
    )

    def __init__(self):
        self.column_identifier = ColumnIdentifier()

    def parse(self, file_path: str) -> ParseResult:
        """
        解析Word文件

        Args:
            file_path: Word文件路径

        Returns:
            ParseResult: 解析结果
        """
        if Document is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

        warnings = []
        items = []

        doc = Document(file_path)

        # 遍历所有表格，解析附表明细
        for table_index, table in enumerate(doc.tables):
            table_items, table_warnings = self._parse_table(table, table_index)
            items.extend(table_items)
            warnings.extend(table_warnings)

        return ParseResult(items, warnings)

    def _parse_reason_section(self, doc) -> tuple:
        """
        从Word文档的"事由"段落提取汇总数据

        Args:
            doc: Word文档对象

        Returns:
            Tuple[List[Dict], List[str]]: (项目列表, 警告列表)
        """
        items = []
        warnings = []
        found = set()  # 记录已处理的单元格文本，避免重复

        # 遍历所有表格查找"事由"单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    # 查找包含"事由"且包含"详见附表"的单元格（汇总数据在事由段落中）
                    if '事由' in cell_text and '详见附表' in cell_text:
                        # 避免重复处理同一行内容
                        text_hash = hash(cell_text[:200])
                        if text_hash in found:
                            continue
                        found.add(text_hash)

                        # 这是事由段落，提取汇总数据
                        summary_items = self._extract_summary_from_text(cell_text)
                        items.extend(summary_items)
                        if summary_items:
                            warnings.append(f"从事由段落提取到 {len(summary_items)} 项汇总数据")

        return items, warnings

    def _extract_summary_from_text(self, text: str) -> List[Dict]:
        """
        从事由文本中提取汇总工程量

        Args:
            text: 事由文本

        Returns:
            List[Dict]: 工程量项目列表
        """
        items = []

        # 定义已知的汇总项模式（从事由段落分析得出）
        summary_patterns = [
            # 4mm²接地跨接黄绿双色线
            (r'4mm[²²]\s*接地跨接黄绿双色线\s*(\d+(?:\.\d+)?)\s*(米|m)', '4mm²接地跨接黄绿双色线', '米'),
            # 三防接线盒
            (r'三防接线盒\s*(\d+(?:\.\d+)?)\s*(个)', '三防接线盒', '个'),
            # 防爆接线盒
            (r'防爆接线盒\s*(\d+(?:\.\d+)?)\s*(个)', '防爆接线盒', '个'),
            # 防爆活接头
            (r'防爆活接头\s*(\d+(?:\.\d+)?)\s*(个)', '防爆活接头', '个'),
            # 10#热镀锌槽钢
            (r'10#\s*热镀锌槽钢\s*(\d+(?:\.\d+)?)\s*(米|m)', '10#热镀锌槽钢', '米'),
            # 热镀锌角钢
            (r'∠?\s*40\s*[x×\*]\s*40\s*[x×\*]\s*4\s*热镀锌角钢\s*(\d+(?:\.\d+)?)\s*(米|m)', '∠40*40*4热镀锌角钢', '米'),
            # 摄像头立柱
            (r'(\d+(?:\.\d+)?)\s*米\s*摄像头立柱\s*(\d+(?:\.\d+)?)\s*(个|套|台)', None, None),  # 需要特殊处理
            # DN80热镀锌钢管
            (r'DN80\s*热镀锌钢管', 'DN80热镀锌钢管', '米'),
            # 200*200*8mm钢板
            (r'200\s*[x×\*]\s*200\s*[x×\*]\s*8\s*mm\s*钢板', '200*200*8mm钢板', '块'),
        ]

        # 直接搜索具体数值
        # 4mm²接地跨接黄绿双色线
        match = re.search(r'4mm[²²]\s*接地跨接黄绿双色线\s*(\d+(?:\.\d+)?)\s*(米|m)', text)
        if match:
            items.append({
                "name": "4mm²接地跨接黄绿双色线",
                "quantity": float(match.group(1)),
                "unit": "米",
                "sheet": "事由汇总",
                "original_unit": "米"
            })

        # 三防接线盒
        match = re.search(r'三防接线盒\s*(\d+(?:\.\d+)?)\s*(?:个)', text)
        if match:
            items.append({
                "name": "三防接线盒",
                "quantity": float(match.group(1)),
                "unit": "个",
                "sheet": "事由汇总",
                "original_unit": "个"
            })

        # 防爆接线盒
        match = re.search(r'防爆接线盒\s*(\d+(?:\.\d+)?)\s*(?:个)', text)
        if match:
            items.append({
                "name": "防爆接线盒",
                "quantity": float(match.group(1)),
                "unit": "个",
                "sheet": "事由汇总",
                "original_unit": "个"
            })

        # 防爆活接头
        match = re.search(r'防爆活接头\s*(\d+(?:\.\d+)?)\s*(?:个)', text)
        if match:
            items.append({
                "name": "防爆活接头",
                "quantity": float(match.group(1)),
                "unit": "个",
                "sheet": "事由汇总",
                "original_unit": "个"
            })

        # 10#热镀锌槽钢
        match = re.search(r'10#\s*热镀锌槽钢\s*(\d+(?:\.\d+)?)\s*(米|m)', text)
        if match:
            items.append({
                "name": "10#热镀锌槽钢",
                "quantity": float(match.group(1)),
                "unit": "米",
                "sheet": "事由汇总",
                "original_unit": "米"
            })

        # ∠40*40*4热镀锌角钢
        match = re.search(r'∠?\s*40\s*[x×\*]\s*40\s*[x×\*]\s*4\s*热镀锌角钢\s*(\d+(?:\.\d+)?)\s*(米|m)', text)
        if match:
            items.append({
                "name": "∠40*40*4热镀锌角钢",
                "quantity": float(match.group(1)),
                "unit": "米",
                "sheet": "事由汇总",
                "original_unit": "米"
            })

        # 摄像头立柱（3.5米、3米、4米、1.5米、1.8米）
        for height in ['3.5米', '3米', '4米', '1.5米', '1.8米']:
            pattern = rf'{height}\s*摄像头立柱\s*(\d+(?:\.\d+)?)\s*(?:个|套|台)'
            match = re.search(pattern, text)
            if match:
                items.append({
                    "name": f"{height}摄像头立柱",
                    "quantity": float(match.group(1)),
                    "unit": "个",
                    "sheet": "事由汇总",
                    "original_unit": "个"
                })

        return items

    def _parse_table(self, table, table_index: int) -> tuple:
        """
        解析Word表格 - 支持标准工程量清单格式

        Args:
            table: Word表格对象
            table_index: 表格索引

        Returns:
            Tuple[List[Dict], List[str]]: (项目列表, 警告列表)
        """
        items = []
        warnings = []

        # 获取表格所有行
        rows = table.rows

        if len(rows) < 2:
            return items, warnings

        # 检查表头
        header_row = rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]

        # 判断表格类型并解析
        header_text = '|'.join(headers).lower()

        if '穿线管规格' in header_text or '跨接线数量' in header_text:
            # 附表1：穿线管和跨接线明细
            items, warns = self._parse_table_pipe(rows, table_index)
            warnings.extend(warns)
        elif '设备数量' in header_text and '型号' in header_text:
            # 附表2：接线盒及活接头明细
            items, warns = self._parse_table_junction_box(rows, table_index)
            warnings.extend(warns)
        elif '槽钢数量' in header_text or '角钢数量' in header_text:
            # 附表3：槽钢、角钢明细
            items, warns = self._parse_table_steel(rows, table_index)
            warnings.extend(warns)
        elif '单体' in header_text and '材质' in header_text and '名称' in header_text and '单位' in header_text and '数量' in header_text:
            # 附表4：监控立柱明细
            items, warns = self._parse_table_pole(rows, table_index)
            warnings.extend(warns)
        else:
            # 通用表格解析：使用列名识别
            items, warns = self._parse_table_generic(rows, table_index)
            warnings.extend(warns)

        return items, warnings

    def _parse_table_pipe(self, rows, table_index: int) -> tuple:
        """解析附表1：穿线管和跨接线明细"""
        items = []
        warnings = []

        headers = [cell.text.strip() for cell in rows[0].cells]

        # 找到各列索引
        # 列: 单体, 穿线管规格, 单位, 穿线管数量, 跨接线数量
        spec_col = next((i for i, h in enumerate(headers) if '规格' in h), None)
        qty_col = next((i for i, h in enumerate(headers) if h == '穿线管数量'), None)
        gnd_qty_col = next((i for i, h in enumerate(headers) if '跨接线' in h), None)
        unit_col = next((i for i, h in enumerate(headers) if h == '单位'), None)

        for row_idx, row in enumerate(rows[1:], start=2):
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not any(cells):
                continue

            # 提取穿线管工程量
            if spec_col is not None and qty_col is not None and qty_col < len(cells):
                spec = cells[spec_col] if spec_col < len(cells) else ''
                qty = cells[qty_col] if qty_col < len(cells) else ''
                unit = cells[unit_col] if unit_col is not None and unit_col < len(cells) else '米'

                if qty and qty.replace('.', '').isdigit():
                    items.append({
                        "name": f"热镀锌钢管 {spec}",
                        "quantity": float(qty),
                        "unit": unit,
                        "sheet": f"附表1-穿线管",
                        "original_unit": unit
                    })

            # 提取跨接线工程量
            if gnd_qty_col is not None and gnd_qty_col < len(cells):
                gnd_qty = cells[gnd_qty_col]
                if gnd_qty and gnd_qty.replace('.', '').isdigit():
                    items.append({
                        "name": "4mm²接地跨接黄绿双色线",
                        "quantity": float(gnd_qty),
                        "unit": "米",
                        "sheet": f"附表1-跨接线",
                        "original_unit": "米"
                    })

        return items, warnings

    def _parse_table_junction_box(self, rows, table_index: int) -> tuple:
        """解析附表2：接线盒及活接头明细"""
        items = []
        warnings = []

        headers = [cell.text.strip() for cell in rows[0].cells]

        # 找到各列索引
        # 列: 序号, 单体, 单位, 设备数量, 名称, 型号, 单位, 数量
        name_col = next((i for i, h in enumerate(headers) if h == '名称'), None)
        model_col = next((i for i, h in enumerate(headers) if h == '型号'), None)
        qty_col = next((i for i, h in enumerate(headers) if h == '数量'), None)
        unit_col = next((i for i, h in enumerate(headers[-1:]+['个']) if h in headers), None)

        for row_idx, row in enumerate(rows[1:], start=2):
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not any(cells):
                continue

            if name_col is not None and name_col < len(cells):
                name = cells[name_col]
                qty = cells[qty_col] if qty_col is not None and qty_col < len(cells) else ''
                unit = '个'

                if qty and qty.replace('.', '').isdigit():
                    items.append({
                        "name": name,
                        "quantity": float(qty),
                        "unit": unit,
                        "sheet": f"附表2-接线盒活接头",
                        "original_unit": unit
                    })

        return items, warnings

    def _parse_table_steel(self, rows, table_index: int) -> tuple:
        """解析附表3：槽钢、角钢明细"""
        items = []
        warnings = []

        # 表格结构：前半部分是槽钢，后半部分是角钢
        # 每部分第一行是表头，后面是数据
        # 列: 单体, 专业, 名称, 数量, 单位, 槽钢/角钢数量, 单位

        steel_type = None  # '槽钢' or '角钢'

        for row_idx, row in enumerate(rows):
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not any(cells):
                continue

            # 检测是否是表头行（包含"槽钢数量"或"角钢数量"）
            header_text = '|'.join(cells)

            if '槽钢数量' in header_text:
                steel_type = '槽钢'
                continue  # 跳过表头行
            elif '角钢数量' in header_text:
                steel_type = '角钢'
                continue  # 跳过表头行

            # 处理数据行
            if steel_type and len(cells) >= 6:
                name = cells[2] if len(cells) > 2 else ''  # 名称列
                qty = cells[5] if len(cells) > 5 else ''  # 钢材数量列

                if qty:
                    try:
                        qty_val = float(qty)
                        if steel_type == '槽钢':
                            items.append({
                                "name": f"10#热镀锌槽钢 ({name})",
                                "quantity": qty_val,
                                "unit": "米",
                                "sheet": "附表3-槽钢",
                                "original_unit": "米"
                            })
                        else:  # 角钢
                            items.append({
                                "name": f"∠40*40*4热镀锌角钢 ({name})",
                                "quantity": qty_val,
                                "unit": "米",
                                "sheet": "附表3-角钢",
                                "original_unit": "米"
                            })
                    except ValueError:
                        pass

        return items, warnings

    def _parse_table_pole(self, rows, table_index: int) -> tuple:
        """解析附表4：监控立柱明细"""
        items = []
        warnings = []

        # 格式: 单体, 专业, 材质, 名称(安装高度X米), 单位, 数量
        # Row0是表头，从Row1开始是数据

        for row_idx, row in enumerate(rows[1:], start=2):
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not any(cells):
                continue

            if len(cells) < 6:
                continue

            # 提取数据
            # cells[2]: 材质 (热镀锌钢管DN80 和 钢板200*200*8mm)
            # cells[3]: 名称 (安装高度3.5米)
            # cells[4]: 单位 (套)
            # cells[5]: 数量

            material = cells[2].replace('|', ' ') if len(cells) > 2 else ''
            name = cells[3].replace('|', ' ') if len(cells) > 3 else ''
            unit = cells[4] if len(cells) > 4 else '套'
            qty_text = cells[5] if len(cells) > 5 else ''

            if qty_text:
                try:
                    qty_val = float(qty_text)
                    # 从名称中提取高度
                    import re
                    match = re.search(r'安装高度(\d+\.?\d*)米', name)
                    height = match.group(1) if match else ''

                    items.append({
                        "name": f"监控立柱 高度{height}米 ({material})",
                        "quantity": qty_val,
                        "unit": unit,
                        "sheet": "附表4-监控立柱",
                        "original_unit": unit
                    })
                except ValueError:
                    pass

        return items, warnings

    def _parse_table_generic(self, rows, table_index: int) -> tuple:
        """通用表格解析：基于列名识别"""
        items = []
        warnings = []

        if len(rows) < 2:
            return items, warnings

        headers = [cell.text.strip() for cell in rows[0].cells]
        column_map = ColumnIdentifier.find_column_index(headers)

        name_col = column_map.get("name")
        quantity_col = column_map.get("quantity")
        unit_col = column_map.get("unit")

        if name_col is None and quantity_col is None:
            # 尝试其他方式识别
            for row_idx, row in enumerate(rows[1:], start=2):
                cells = [cell.text.strip() for cell in row.cells]
                if not cells or not any(cells):
                    continue

                # 检查是否包含"共计"模式
                has_qijei = any('共计' in cell for cell in cells)
                if has_qijei:
                    for cell_text in cells:
                        if '共计' not in cell_text:
                            continue
                        match = self.QIJEI_PATTERN.search(cell_text)
                        if match:
                            quantity = float(match.group(1))
                            unit = match.group(2)
                            name = self.QIJEI_PATTERN.sub("", cell_text).strip()
                            name = ' '.join(name.split())
                            if name:
                                items.append({
                                    "name": name,
                                    "quantity": quantity,
                                    "unit": unit,
                                    "sheet": f"表格{table_index + 1}",
                                    "original_unit": unit
                                })

            items = self._deduplicate_items(items)
            return items, warnings

        for row_idx, row in enumerate(rows[1:], start=2):
            cells = [cell.text.strip() for cell in row.cells]

            if name_col is not None and quantity_col is not None:
                if name_col < len(cells) and quantity_col < len(cells):
                    name = cells[name_col]
                    qty_text = cells[quantity_col]
                    unit = cells[unit_col] if unit_col is not None and unit_col < len(cells) else "项"

                    item = self._parse_normal_row(name, qty_text, unit, table_index, row_idx)
                    if item:
                        items.append(item)

        items = self._deduplicate_items(items)
        return items, warnings

    def _deduplicate_items(self, items: List[Dict]) -> List[Dict]:
        """去重：如果多个项目名称和工程量相同，只保留一个"""
        seen = set()
        unique_items = []
        for item in items:
            key = (item['name'], item['quantity'], item['unit'])
            if key not in seen:
                seen.add(key)
                unique_items.append(item)
        return unique_items

    def _parse_normal_row(self, name: str, qty_text: str, unit: str,
                         table_index: int, row_index: int) -> Optional[Dict]:
        """
        解析常规格式的行

        Args:
            name: 项目名称
            qty_text: 数量文本
            unit: 单位
            table_index: 表格索引
            row_index: 行索引

        Returns:
            Dict: 项目数据
        """
        if not name or not qty_text:
            return None

        try:
            # 清理数量文本
            clean_qty = qty_text.replace("≈", "").replace("约", "").strip()
            quantity = float(clean_qty)
        except ValueError:
            # 尝试查找"共计"模式
            match = self.QIJEI_PATTERN.search(qty_text)
            if match:
                quantity = float(match.group(1))
                unit = match.group(2)
            else:
                return None

        return {
            "name": name.strip(),
            "quantity": quantity,
            "unit": unit.strip() or "项",
            "sheet": f"表格{table_index + 1}",
            "original_unit": unit
        }


class FileParser:
    """文件解析器（自动识别类型）"""

    PARSER_MAP = {
        ".xlsx": ExcelParser,
        ".xls": ExcelParser,
        ".docx": WordParser,
    }

    def __init__(self):
        self.parsers = {ext: parser_class() for ext, parser_class in self.PARSER_MAP.items()}

    def parse(self, file_path: str) -> ParseResult:
        """
        自动识别文件类型并解析

        Args:
            file_path: 文件路径

        Returns:
            ParseResult: 解析结果
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        parser = self.parsers.get(ext)

        if parser is None:
            raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {list(self.PARSER_MAP.keys())}")

        return parser.parse(file_path)


if __name__ == "__main__":
    # 测试代码
    parser = FileParser()

    # 测试Excel解析
    # result = parser.parse("test.xlsx")

    # 测试Word解析
    # result = parser.parse("test.docx")

    print("FileParser 已初始化，支持格式:", list(FileParser.PARSER_MAP.keys()))
